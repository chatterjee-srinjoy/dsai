"""Build the Homework 2 .docx submission file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

REPO = "https://github.com/chatterjee-srinjoy/dsai/blob/main"

# ── Helper functions ──────────────────────────────────────

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return h

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def link_bullet(label, path):
    url = f"{REPO}/{path}"
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(label)
    run.bold = True
    p.add_run(f": {url}")
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    return table

def spacer():
    doc.add_paragraph()

# ── Title ─────────────────────────────────────────────────

title = doc.add_heading("Homework 2: AI Agent System with RAG and Tools", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

para("Srinjoy Chatterjee", bold=True)
para("FDA Device Recall AI Agent System", italic=True)
spacer()

# ══════════════════════════════════════════════════════════
# SECTION 1: WRITING COMPONENT (25 pts)
# ══════════════════════════════════════════════════════════

heading("1. Writing Component", level=1)

doc.add_paragraph(
    "This project builds an AI agent system that analyzes FDA medical device "
    "recalls by combining three techniques we learned over the past three weeks: "
    "multi-agent orchestration, retrieval-augmented generation (RAG), and function "
    "calling with tools. The system is built entirely in Python and uses a local "
    "Ollama LLM (smollm2:1.7b) for all agent interactions."
)

doc.add_paragraph(
    "The system works as a three-agent pipeline. Agent 1 (the Data Fetcher) uses "
    "function calling to query the openFDA Device Recall API and retrieve live "
    "recall records for a given year. The LLM decides the tool arguments based on "
    "the task prompt, calls the get_fda_recalls() function, and returns a structured "
    "DataFrame of results. Agent 2 (the Context Researcher) uses RAG to search a "
    "local knowledge base I created — a text file containing FDA domain expertise "
    "about recall classification levels, common root causes, the regulatory process, "
    "and patient safety impact. This context gives the system understanding that the "
    "raw API data alone cannot provide. Agent 3 (the Executive Reporter) receives "
    "both the live data from Agent 1 and the synthesized context from Agent 2, then "
    "writes a professional executive brief with sections for overview, key findings, "
    "regulatory context, and recommendations."
)

doc.add_paragraph(
    "The main design challenge was making the three components work together "
    "coherently. The function calling component required injecting the tool function "
    "into the functions module so Ollama's tool dispatch could find it. For RAG, I "
    "chose a text-based knowledge base over CSV or SQLite because the domain context "
    "is naturally unstructured (paragraphs of regulatory explanations), and a simple "
    "line-based substring search was sufficient given the focused scope of the "
    "knowledge base. The multi-agent chaining uses the agent_run() helper from the "
    "course's functions.py, passing each agent's output as the next agent's task "
    "input. One limitation is that smollm2:1.7b is a small model, so its outputs "
    "are sometimes imprecise — a larger model would produce better reports, but the "
    "pipeline architecture itself is solid and model-agnostic."
)

spacer()

# ══════════════════════════════════════════════════════════
# SECTION 2: GIT REPOSITORY LINKS (25 pts)
# ══════════════════════════════════════════════════════════

heading("2. Code — Git Repository Links", level=1)

para("All code is in the dsai repository on GitHub:", bold=False)

link_bullet("Multi-agent orchestration script (Lab 1)",
            "06_agents/lab_prompt_design.py")

link_bullet("RAG implementation (Lab 2)",
            "07_rag/lab_custom_rag_query.py")

link_bullet("RAG knowledge base data",
            "07_rag/data/fda_recall_knowledge.txt")

link_bullet("Function calling / tool definitions (Lab 3)",
            "08_function_calling/lab_multi_agent_with_tools.py")

link_bullet("Main system file (Homework 2 — all components combined)",
            "08_function_calling/homework2_fda_agent.py")

link_bullet("Documentation",
            "08_function_calling/DOCUMENTATION.md")

link_bullet("Helper functions (course-provided, used by all scripts)",
            "08_function_calling/functions.py")

spacer()

# ══════════════════════════════════════════════════════════
# SECTION 3: SCREENSHOTS / OUTPUTS (25 pts)
# ══════════════════════════════════════════════════════════

heading("3. Screenshots / Outputs", level=1)

para(
    "Note: Insert screenshots from running each script in the terminal. "
    "Below are sample outputs from each component.",
    italic=True,
)
spacer()

heading("3.1 Multi-Agent Workflow (lab_prompt_design.py)", level=2)
para("This screenshot shows three agents chaining together: "
     "Agent 1 summarizes FDA recall data, Agent 2 analyzes trends, "
     "and Agent 3 writes an executive report.", italic=True)
para("[INSERT SCREENSHOT: Run 'cd 06_agents && python3 lab_prompt_design.py']",
     bold=True)
spacer()

heading("3.2 RAG Retrieval and Response (lab_custom_rag_query.py)", level=2)
para("This screenshot shows the RAG workflow: searching the knowledge base "
     "for matching lines, then passing retrieved context to the LLM for "
     "context-aware answers.", italic=True)
para("[INSERT SCREENSHOT: Run 'cd 07_rag && python3 lab_custom_rag_query.py']",
     bold=True)
spacer()

heading("3.3 Function Calling / Tool Usage (lab_multi_agent_with_tools.py)", level=2)
para("This screenshot shows Agent 1 calling the get_fda_recalls() tool to "
     "fetch live data from the openFDA API, returning a DataFrame of recall "
     "records. Agent 2 then analyzes the data.", italic=True)
para("[INSERT SCREENSHOT: Run 'cd 08_function_calling && python3 lab_multi_agent_with_tools.py']",
     bold=True)
spacer()

heading("3.4 Combined System (homework2_fda_agent.py)", level=2)
para("This screenshot shows the unified system with all three components: "
     "function calling (Agent 1), RAG (Agent 2), and the final executive "
     "report (Agent 3).", italic=True)
para("[INSERT SCREENSHOT: Run 'cd 08_function_calling && python3 homework2_fda_agent.py']",
     bold=True)
spacer()

# ══════════════════════════════════════════════════════════
# SECTION 4: DOCUMENTATION (25 pts)
# ══════════════════════════════════════════════════════════

heading("4. Documentation", level=1)

# 4.1 System Architecture
heading("4.1 System Architecture", level=2)

para("The system uses a 3-agent pipeline where each agent has a specific role:")

add_table(
    ["Agent", "Role", "Component", "Input", "Output"],
    [
        ["Agent 1: Data Fetcher",
         "Queries openFDA API",
         "Function Calling",
         "LLM picks tool args (year, limit)",
         "DataFrame of recall records"],
        ["Agent 2: Context Researcher",
         "Searches local knowledge base",
         "RAG",
         "search results for root causes, classifications, safety",
         "Synthesized domain context"],
        ["Agent 3: Executive Reporter",
         "Writes professional report",
         "Multi-Agent",
         "Data from Agent 1 + context from Agent 2",
         "Executive brief (Overview, Findings, Context, Recommendations)"],
    ],
)
spacer()

para("Workflow: Agent 1 output and Agent 2 output are both passed to Agent 3, "
     "which synthesizes them into a final executive report.")
spacer()

# 4.2 RAG Data Source
heading("4.2 RAG Data Source", level=2)

bullet("File: ", bold_prefix="")
para("07_rag/data/fda_recall_knowledge.txt")
bullet("Type: ", bold_prefix="")
para("Plain text file (~40 lines) of FDA domain knowledge")
bullet("Search method: ", bold_prefix="")
para("Case-insensitive substring matching across all lines")

para("The knowledge base contains information not available from the API:")
add_table(
    ["Topic", "Content"],
    [
        ["Recall Classifications", "Class I / II / III definitions, severity, examples, required actions"],
        ["Root Causes", "Process Control, Nonconforming Material, Software Design, Design, Labeling, Under Investigation"],
        ["Regulatory Process", "CDRH oversight, 21 CFR Part 806, recall status types (Ongoing, Terminated, Completed)"],
        ["Metrics & Trends", "Annual recall volumes (~3,000-4,000/year), top firms, product categories"],
        ["Patient Safety", "Class I immediate action requirements, MAUDE database, facility quarantine processes"],
    ],
)
spacer()

# 4.3 Tool Functions
heading("4.3 Tool Functions", level=2)

add_table(
    ["Function", "Purpose", "Parameters", "Returns"],
    [
        ["get_fda_recalls(year, limit)",
         "Query openFDA Device Recall API for recall records",
         "year (int): e.g. 2024; limit (int): 1-1000",
         "pandas DataFrame with columns: recall_number, date, firm, root_cause, product_code, status"],
        ["search_fda_knowledge(query, path)",
         "Search local knowledge base for matching lines",
         "query (str): search term; path (str): file path",
         "String of matching lines concatenated"],
    ],
)
spacer()

# 4.4 Technical Details
heading("4.4 Technical Details", level=2)

para("API:", bold=True)
bullet("Endpoint: https://api.fda.gov/device/recall.json")
bullet("Auth: Optional API_KEY in .env (higher rate limits)")
bullet("Query: Lucene search on event_date_initiated with date range filter")

para("LLM:", bold=True)
bullet("Provider: Ollama (local, http://localhost:11434)")
bullet("Model: smollm2:1.7b")
bullet("Tools API: /api/chat with tools parameter for function calling")

para("Python Packages:", bold=True)
add_table(
    ["Package", "Version", "Purpose"],
    [
        ["requests", "2.32+", "HTTP requests to FDA API and Ollama"],
        ["pandas", "3.0+", "Data manipulation and DataFrames"],
        ["python-dotenv", "1.2+", "Loading .env for API keys"],
        ["tabulate", "0.10+", "df.to_markdown() for table output"],
    ],
)
spacer()

para("File Structure:", bold=True)
doc.add_paragraph(
    "dsai/\n"
    "├── .env                              # API keys (gitignored)\n"
    "├── 06_agents/\n"
    "│   ├── functions.py                  # Agent helpers (course-provided)\n"
    "│   └── lab_prompt_design.py          # Lab 1: Multi-agent prompt design\n"
    "├── 07_rag/\n"
    "│   ├── functions.py                  # RAG helpers (course-provided)\n"
    "│   ├── data/\n"
    "│   │   └── fda_recall_knowledge.txt  # RAG knowledge base\n"
    "│   └── lab_custom_rag_query.py       # Lab 2: RAG query workflow\n"
    "└── 08_function_calling/\n"
    "    ├── functions.py                  # Function calling helpers (course-provided)\n"
    "    ├── lab_multi_agent_with_tools.py # Lab 3: Multi-agent with FDA tool\n"
    "    └── homework2_fda_agent.py        # Main system (all 3 components)\n",
    style="No Spacing",
)
spacer()

# 4.5 Usage Instructions
heading("4.5 Usage Instructions", level=2)

para("Step 1: Install dependencies", bold=True)
doc.add_paragraph("pip install requests pandas python-dotenv tabulate",
                   style="No Spacing")
spacer()

para("Step 2: Configure API key (optional)", bold=True)
doc.add_paragraph("Create .env in the repo root with: API_KEY=your_fda_key",
                   style="No Spacing")
spacer()

para("Step 3: Start Ollama", bold=True)
doc.add_paragraph("ollama serve", style="No Spacing")
doc.add_paragraph("ollama pull smollm2:1.7b", style="No Spacing")
spacer()

para("Step 4: Run the main system", bold=True)
doc.add_paragraph("cd 08_function_calling", style="No Spacing")
doc.add_paragraph("python3 homework2_fda_agent.py", style="No Spacing")
spacer()

para("Step 5: Run individual lab scripts (optional)", bold=True)
doc.add_paragraph(
    "cd 06_agents && python3 lab_prompt_design.py      # Lab 1\n"
    "cd 07_rag && python3 lab_custom_rag_query.py       # Lab 2\n"
    "cd 08_function_calling && python3 lab_multi_agent_with_tools.py  # Lab 3",
    style="No Spacing",
)

# ── Save ──────────────────────────────────────────────────

out_path = os.path.join(os.path.dirname(__file__), "Homework2_Srinjoy_Chatterjee.docx")
doc.save(out_path)
print(f"✅ Saved: {out_path}")
